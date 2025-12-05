"""
将图片中的判断题识别并直接生成 Word 文档

参考 image_to_word.py 的结构，专门用于识别判断题
"""

import argparse
import base64
import json
import os
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openai import OpenAI
from pydantic import BaseModel


# ==================== 数据模型 ====================

class TrueFalseItem(BaseModel):
    """判断题数据模型"""
    title: str  # 题目内容


class TrueFalseList(BaseModel):
    """判断题列表"""
    questions: list[TrueFalseItem]


# ==================== 豆包模型配置 ====================

DOUBAO_API_KEY = os.getenv("DOUBAO_API_KEY", "a871c6fc-c014-4451-9db5-125400419b25")
DOUBAO_BASE_URL = "https://ark.cn-beijing.volces.com/api/v3"
DOUBAO_MODEL = os.getenv("DOUBAO_MODEL", "doubao-seed-1-6-251015")

SYSTEM_PROMPT = """你是一个专业的题目识别助手。请仔细分析图片中的判断题，识别出所有题目。

要求：
1. 识别图片中所有的判断题
2. 提取每道判断题的题目内容（title）
3. 判断题通常是一个陈述句，需要判断其正确或错误
4. 请按照图片中题目出现的顺序提取
5. 不要提取题目的序号，只提取题目内容本身
"""


# ==================== 图片识别功能 ====================

def encode_image_to_base64(image_path: str) -> str:
    """将图片编码为 base64 字符串"""
    with open(image_path, "rb") as image_file:
        return base64.standard_b64encode(image_file.read()).decode("utf-8")


def get_image_mime_type(image_path: str) -> str:
    """根据文件扩展名获取 MIME 类型"""
    suffix = Path(image_path).suffix.lower()
    mime_types = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".bmp": "image/bmp",
    }
    return mime_types.get(suffix, "image/jpeg")


def recognize_truefalse_from_images(image_paths: list[str]) -> list[dict]:
    """
    使用豆包模型识别图片中的判断题
    
    Args:
        image_paths: 图片文件路径列表
        
    Returns:
        识别出的判断题列表
    """
    if not DOUBAO_API_KEY:
        raise ValueError("请设置环境变量 DOUBAO_API_KEY")
    
    client = OpenAI(
        api_key=DOUBAO_API_KEY,
        base_url=DOUBAO_BASE_URL,
    )
    
    # 构建包含多张图片的消息内容
    content = []
    
    # 添加文字提示
    content.append({
        "type": "text",
        "text": "请识别以下图片中的所有判断题。"
    })
    
    # 添加所有图片
    valid_images = 0
    for image_path in image_paths:
        if not os.path.exists(image_path):
            print(f"警告: 图片文件不存在: {image_path}")
            continue
            
        base64_image = encode_image_to_base64(image_path)
        mime_type = get_image_mime_type(image_path)
        
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:{mime_type};base64,{base64_image}"
            }
        })
        valid_images += 1
    
    if valid_images == 0:
        raise ValueError("没有有效的图片文件")
    
    print(f"正在识别 {valid_images} 张图片中的判断题...")
    
    # 使用 Chat Completions API 的结构化输出
    response = client.beta.chat.completions.parse(
        model=DOUBAO_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": content}
        ],
        temperature=0.1,
        response_format=TrueFalseList,
    )
    
    # 解析响应
    result = response.choices[0].message.parsed
    
    # 转换为字典列表格式
    questions = [
        {"title": q.title}
        for q in result.questions
    ]
    
    print(f"成功识别出 {len(questions)} 道判断题")
    return questions


# ==================== Word 文档生成功能 ====================

def remove_table_borders(document: Document) -> None:
    """移除文档中所有表格的边框"""
    for table in document.tables:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "nil")
            borders.append(border)
        tbl_pr.append(borders)


def apply_yahei_font(run, size_pt: float) -> None:
    """应用微软雅黑字体"""
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size_pt)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr, font in (
        ("w:ascii", "Microsoft YaHei"),
        ("w:hAnsi", "Microsoft YaHei"),
        ("w:cs", "Microsoft YaHei"),
        ("w:eastAsia", "微软雅黑"),
    ):
        r_fonts.set(qn(attr), font)


def format_cell_text(cell, font_size: float) -> None:
    """格式化单元格文本"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            apply_yahei_font(run, font_size)


def format_paragraph_text(paragraph, font_size: float) -> None:
    """格式化段落文本"""
    for run in paragraph.runs:
        apply_yahei_font(run, font_size)


def strip_leading_number(title: str) -> str:
    """移除题目前的序号"""
    return re.sub(r"^\s*\d+[\.)、．]\s*", "", title)


def build_truefalse_questions(items: Iterable[dict], document: Document, start_number: int = 1) -> int:
    """
    为每道判断题创建带编号的段落
    
    Args:
        items: 判断题列表
        document: Word 文档对象
        start_number: 起始编号
        
    Returns:
        下一个可用的编号
    """
    current_number = start_number
    
    for item in items:
        # 创建带编号的段落
        paragraph = document.add_paragraph()
        paragraph.style = "List Number"
        
        # 添加题目内容
        title_text = strip_leading_number(item["title"].strip())
        run = paragraph.add_run(f"{title_text}（    ）")
        apply_yahei_font(run, 12)
        
        current_number += 1
    
    return current_number


def build_truefalse_tables(items: Iterable[dict], document: Document) -> None:
    """
    为每道判断题创建表格（1行1列，用于格式一致性）
    
    Args:
        items: 判断题列表
        document: Word 文档对象
    """
    for item in items:
        table = document.add_table(rows=1, cols=1)
        table.allow_autofit = False

        # 题目单元格
        title_cell = table.cell(0, 0)
        title_paragraph = title_cell.paragraphs[0]
        title_paragraph.style = "List Number"
        
        # 添加题目内容，末尾添加判断题答案括号
        title_text = strip_leading_number(item["title"].strip())
        title_paragraph.text = f"{title_text}（    ）"
        format_cell_text(title_cell, 12)

        document.add_paragraph()


def get_existing_question_count(document: Document) -> int:
    """获取文档中已有的题目数量"""
    count = 0
    for paragraph in document.paragraphs:
        if paragraph.style and paragraph.style.name == "List Number":
            count += 1
    return count


def generate_word_document(questions: list[dict], output_path: Path, use_table: bool = False) -> None:
    """
    生成 Word 文档
    
    Args:
        questions: 判断题列表
        output_path: 输出文件路径
        use_table: 是否使用表格格式（默认使用段落格式）
    """
    document = Document(output_path) if output_path.exists() else Document()
    
    if use_table:
        build_truefalse_tables(questions, document)
        remove_table_borders(document)
    else:
        # 获取已有题目数量，用于确定起始编号
        start_number = get_existing_question_count(document) + 1
        build_truefalse_questions(questions, document, start_number)
    
    document.save(output_path)


# ==================== 主程序 ====================

def main():
    parser = argparse.ArgumentParser(
        description="将图片中的判断题识别并生成 Word 文档"
    )
    parser.add_argument(
        "images",
        nargs="+",
        help="要识别的图片文件路径（支持多张图片）"
    )
    parser.add_argument(
        "-o", "--output",
        type=Path,
        default=Path("truefalse_questions.docx"),
        help="输出 Word 文档路径（默认: truefalse_questions.docx）"
    )
    parser.add_argument(
        "--save-json",
        type=Path,
        help="同时保存 JSON 格式的识别结果"
    )
    parser.add_argument(
        "--append",
        action="store_true",
        help="追加到现有文档而不是创建新文档"
    )
    parser.add_argument(
        "--table",
        action="store_true",
        help="使用表格格式输出（默认使用段落格式）"
    )
    
    args = parser.parse_args()
    
    # 确保输出文件有 .docx 扩展名
    output_path = args.output
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")
        print(f"输出文件已重命名为: {output_path}")
    
    try:
        # 步骤1: 识别图片中的判断题
        questions = recognize_truefalse_from_images(args.images)
        
        if not questions:
            print("未识别到任何判断题")
            return 1
        
        # 步骤2: 保存 JSON（可选）
        if args.save_json:
            json_output = json.dumps(questions, ensure_ascii=False, indent=2)
            args.save_json.write_text(json_output, encoding="utf-8")
            print(f"JSON 结果已保存到: {args.save_json}")
        
        # 步骤3: 生成 Word 文档
        if not args.append and output_path.exists():
            output_path.unlink()  # 删除现有文件以创建新文档
            
        generate_word_document(questions, output_path, use_table=args.table)
        print(f"Word 文档已生成: {output_path}")
        
    except Exception as e:
        print(f"错误: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())
