
import json
import pytest
from docx import Document
from src.tools.word_generator import save_questions_word, generate_word_document

class TestWordGeneratorAppend:
    """Tests for appending questions to Word documents with correct section placement."""

    @pytest.fixture
    def mixed_questions(self):
        return {
            "multiple_choice": [
                {"title": "MC Q1", "options": {"a": "1", "b": "2", "c": "3", "d": "4"}}
            ],
            "true_false": [
                {"title": "TF Q1"}
            ]
        }

    def test_append_mc_to_mixed_doc(self, tmp_path, mixed_questions):
        """Test appending a multiple choice question to a document that already has both types."""
        output_path = tmp_path / "test_append_mc.docx"
        
        # 1. Create initial document
        save_questions_word.invoke({
            "questions_json": json.dumps(mixed_questions),
            "output_path": str(output_path),
            "question_type": "auto"
        })
        
        # 2. Append new MC question
        new_mc = [{"title": "MC Q2", "options": {"a": "5", "b": "6", "c": "7", "d": "8"}}]
        save_questions_word.invoke({
            "questions_json": json.dumps(new_mc),
            "output_path": str(output_path),
            "question_type": "multiple_choice",
            "append": True
        })
        
        # 3. Verify order
        doc = Document(output_path)
        
        # We expect:
        # Header: 一、选择题
        # Table: MC Q1
        # Table: MC Q2  <-- Inserted here
        # Header: 二、判断题
        # Table: TF Q1
        
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "一、选择题" in paragraphs
        assert "二、判断题" in paragraphs
        
        # Check tables order
        tables = doc.tables
        assert len(tables) == 3
        
        # Helper to get text from first cell of table
        def get_table_text(table):
            return table.cell(0, 0).text
            
        assert "MC Q1" in get_table_text(tables[0])
        assert "MC Q2" in get_table_text(tables[1])
        assert "TF Q1" in get_table_text(tables[2])

    def test_append_tf_to_mixed_doc(self, tmp_path, mixed_questions):
        """Test appending a true/false question to a document that already has both types."""
        output_path = tmp_path / "test_append_tf.docx"
        
        # 1. Create initial document
        save_questions_word.invoke({
            "questions_json": json.dumps(mixed_questions),
            "output_path": str(output_path),
            "question_type": "auto"
        })
        
        # 2. Append new TF question
        new_tf = [{"title": "TF Q2"}]
        save_questions_word.invoke({
            "questions_json": json.dumps(new_tf),
            "output_path": str(output_path),
            "question_type": "true_false",
            "append": True
        })
        
        # 3. Verify order
        doc = Document(output_path)
        
        tables = doc.tables
        assert len(tables) == 3
        
        def get_table_text(table):
            return table.cell(0, 0).text
            
        assert "MC Q1" in get_table_text(tables[0])
        assert "TF Q1" in get_table_text(tables[1])
        assert "TF Q2" in get_table_text(tables[2])

    def test_insert_mc_section_before_tf(self, tmp_path):
        """Test inserting MC section when only TF section exists."""
        output_path = tmp_path / "test_insert_mc.docx"
        
        # 1. Create document with only TF
        tf_questions = [{"title": "TF Q1"}]
        save_questions_word.invoke({
            "questions_json": json.dumps(tf_questions),
            "output_path": str(output_path),
            "question_type": "true_false"
        })
        
        # 2. Append MC question
        new_mc = [{"title": "MC Q1", "options": {"a": "1", "b": "2", "c": "3", "d": "4"}}]
        save_questions_word.invoke({
            "questions_json": json.dumps(new_mc),
            "output_path": str(output_path),
            "question_type": "multiple_choice",
            "append": True
        })
        
        # 3. Verify order
        doc = Document(output_path)
        
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Should have both headers now
        assert "一、选择题" in paragraphs
        assert "二、判断题" in paragraphs
        
        # MC header should come before TF header
        mc_idx = paragraphs.index("一、选择题")
        tf_idx = paragraphs.index("二、判断题")
        assert mc_idx < tf_idx
        
        tables = doc.tables
        assert len(tables) == 2
        
        def get_table_text(table):
            return table.cell(0, 0).text
            
        # MC table should be first
        assert "MC Q1" in get_table_text(tables[0])
        assert "TF Q1" in get_table_text(tables[1])

    def test_append_tf_creates_section(self, tmp_path):
        """Test appending TF creates section if it doesn't exist."""
        output_path = tmp_path / "test_append_tf_section.docx"
        
        # 1. Create document with only MC
        mc_questions = [{"title": "MC Q1", "options": {"a": "1", "b": "2", "c": "3", "d": "4"}}]
        save_questions_word.invoke({
            "questions_json": json.dumps(mc_questions),
            "output_path": str(output_path),
            "question_type": "multiple_choice"
        })
        
        # 2. Append TF question
        new_tf = [{"title": "TF Q1"}]
        save_questions_word.invoke({
            "questions_json": json.dumps(new_tf),
            "output_path": str(output_path),
            "question_type": "true_false",
            "append": True
        })
        
        # 3. Verify order
        doc = Document(output_path)
        
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "一、选择题" in paragraphs
        assert "二、判断题" in paragraphs
        
        tables = doc.tables
        assert len(tables) == 2
        
        def get_table_text(table):
            return table.cell(0, 0).text
            
        assert "MC Q1" in get_table_text(tables[0])
        assert "TF Q1" in get_table_text(tables[1])
