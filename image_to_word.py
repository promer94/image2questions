"""
将图片中的选择题识别并直接生成 Word 文档

结合 image_to_json.py 和 json_to_word.py 的功能
"""

import argparse
import base64
import json
import os
import re
from pathlib import Path
from typing import Iterable, Sequence

from dotenv import load_dotenv
from docx import Document

load_dotenv()
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openai import OpenAI
from pydantic import BaseModel


# ==================== 数据模型 ====================

class Options(BaseModel):
    a: str
    b: str
    c: str
    d: str


class QuestionItem(BaseModel):
    title: str
    options: Options


class QuestionList(BaseModel):
    questions: list[QuestionItem]


# ==================== 豆包模型配置 ====================

DOUBAO_API_KEY = os.getenv("DOUBAO_API_KEY", "")
DOUBAO_BASE_URL = "https://ark.cn-beijing.volces.com/api/v3"
DOUBAO_MODEL = os.getenv("DOUBAO_MODEL", "doubao-seed-1-6-251015")

SYSTEM_PROMPT = """你是一个专业的题目识别助手。请仔细分析图片中的选择题，识别出所有题目及其选项。

要求：
1. 识别图片中所有的选择题
2. 提取每道题的题目内容（title）和四个选项（a、b、c、d）
3. 如果某个选项不存在，对应的值设为空字符串
"""


# ==================== 图片识别功能 ====================

def encode_image_to_base64(image_path: str) -> str:
    """将图片编码为 base64 字符串"""
    with open(image_path, "rb") as image_file:
        return base64.standard_b64encode(image_file.read()).decode("utf-8")


def get_image_mime_type(image_path: str) -> str:
    """根据文件扩展名获取 MIME 类型"""
    suffix = Path(image_path).suffix.lower()
    mime_types = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".bmp": "image/bmp",
    }
    return mime_types.get(suffix, "image/jpeg")


def recognize_questions_from_images(image_paths: list[str]) -> list[dict]:
    """
    使用豆包模型识别图片中的选择题
    
    Args:
        image_paths: 图片文件路径列表
        
    Returns:
        识别出的选择题列表
    """
    if not DOUBAO_API_KEY:
        raise ValueError("请设置环境变量 DOUBAO_API_KEY")
    
    client = OpenAI(
        api_key=DOUBAO_API_KEY,
        base_url=DOUBAO_BASE_URL,
    )
    
    # 构建包含多张图片的消息内容
    content = []
    
    # 添加文字提示
    content.append({
        "type": "text",
        "text": "请识别以下图片中的所有选择题。"
    })
    
    # 添加所有图片
    valid_images = 0
    for image_path in image_paths:
        if not os.path.exists(image_path):
            print(f"警告: 图片文件不存在: {image_path}")
            continue
            
        base64_image = encode_image_to_base64(image_path)
        mime_type = get_image_mime_type(image_path)
        
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:{mime_type};base64,{base64_image}"
            }
        })
        valid_images += 1
    
    if valid_images == 0:
        raise ValueError("没有有效的图片文件")
    
    print(f"正在识别 {valid_images} 张图片中的选择题...")
    
    # 使用 Chat Completions API 的结构化输出
    response = client.beta.chat.completions.parse(
        model=DOUBAO_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": content}
        ],
        temperature=0.1,
        response_format=QuestionList,
    )
    
    # 解析响应
    result = response.choices[0].message.parsed
    
    # 转换为字典列表格式
    questions = [
        {
            "title": q.title,
            "options": {
                "a": q.options.a,
                "b": q.options.b,
                "c": q.options.c,
                "d": q.options.d,
            }
        }
        for q in result.questions
    ]
    
    print(f"成功识别出 {len(questions)} 道选择题")
    return questions


# ==================== Word 文档生成功能 ====================

def remove_table_borders(document: Document) -> None:
    """移除文档中所有表格的边框"""
    for table in document.tables:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "nil")
            borders.append(border)
        tbl_pr.append(borders)


def apply_yahei_font(run, size_pt: float) -> None:
    """应用微软雅黑字体"""
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size_pt)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr, font in (
        ("w:ascii", "Microsoft YaHei"),
        ("w:hAnsi", "Microsoft YaHei"),
        ("w:cs", "Microsoft YaHei"),
        ("w:eastAsia", "微软雅黑"),
    ):
        r_fonts.set(qn(attr), font)


def format_cell_text(cell, font_size: float) -> None:
    """格式化单元格文本"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            apply_yahei_font(run, font_size)


def strip_leading_number(title: str) -> str:
    """移除题目前的序号"""
    return re.sub(r"^\s*\d+[\.)]\s*", "", title)


def build_tables(items: Iterable[dict], document: Document) -> None:
    """为每道题目创建表格"""
    for item in items:
        table = document.add_table(rows=3, cols=2)
        table.allow_autofit = False

        # 合并第一行作为题目
        title_cell = table.cell(0, 0)
        title_cell.merge(table.cell(0, 1))
        title_paragraph = title_cell.paragraphs[0]
        title_paragraph.style = "List Number"
        title_paragraph.text = strip_leading_number(item["title"].strip())
        format_cell_text(title_cell, 12)

        # 填充选项
        options = item["options"]
        letters: Sequence[str] = ["A", "B", "C", "D"]
        coords = [(1, 0), (1, 1), (2, 0), (2, 1)]

        for letter, coordinate in zip(letters, coords):
            cell = table.cell(*coordinate)
            value = options[letter.lower()].strip()
            cell.text = f"{letter}. {value}"
            format_cell_text(cell, 10)

        document.add_paragraph()


def generate_word_document(questions: list[dict], output_path: Path) -> None:
    """生成 Word 文档"""
    document = Document(output_path) if output_path.exists() else Document()
    build_tables(questions, document)
    remove_table_borders(document)
    document.save(output_path)


# ==================== 主程序 ====================

def main():
    parser = argparse.ArgumentParser(
        description="将图片中的选择题识别并生成 Word 文档"
    )
    parser.add_argument(
        "images",
        nargs="+",
        help="要识别的图片文件路径（支持多张图片）"
    )
    parser.add_argument(
        "-o", "--output",
        type=Path,
        default=Path("multiple-choice.docx"),
        help="输出 Word 文档路径（默认: multiple-choice.docx）"
    )
    parser.add_argument(
        "--save-json",
        type=Path,
        help="同时保存 JSON 格式的识别结果"
    )
    parser.add_argument(
        "--append",
        action="store_true",
        help="追加到现有文档而不是创建新文档"
    )
    
    args = parser.parse_args()
    
    # 确保输出文件有 .docx 扩展名
    output_path = args.output
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")
        print(f"输出文件已重命名为: {output_path}")
    
    try:
        # 步骤1: 识别图片中的选择题
        questions = recognize_questions_from_images(args.images)
        
        if not questions:
            print("未识别到任何选择题")
            return 1
        
        # 步骤2: 保存 JSON（可选）
        if args.save_json:
            json_output = json.dumps(questions, ensure_ascii=False, indent=2)
            args.save_json.write_text(json_output, encoding="utf-8")
            print(f"JSON 结果已保存到: {args.save_json}")
        
        # 步骤3: 生成 Word 文档
        if not args.append and output_path.exists():
            output_path.unlink()  # 删除现有文件以创建新文档
            
        generate_word_document(questions, output_path)
        print(f"Word 文档已生成: {output_path}")
        
    except Exception as e:
        print(f"错误: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())
