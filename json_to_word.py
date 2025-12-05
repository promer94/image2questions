"""Generate a Word document with table cards for each question in JSON data."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
import re
from typing import Iterable, Sequence

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def load_data(path: Path) -> list[dict]:
    """Read and validate the JSON data from a file."""

    raw = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(raw, list):
        raise ValueError("JSON data must be an array of items")

    cleaned: list[dict] = []
    for index, item in enumerate(raw, start=1):
        if not isinstance(item, dict):
            raise ValueError(f"item {index} must be an object")

        title = item.get("title")
        options = item.get("options")
        if not isinstance(title, str) or not isinstance(options, dict):
            raise ValueError(
                f"item {index} must contain 'title' (string) and 'options' (object)"
            )

        if set(options) != {"a", "b", "c", "d"}:
            raise ValueError(f"item {index} must have exactly a/b/c/d keys")

        cleaned.append({"title": title, "options": options})

    return cleaned


def remove_table_borders(document: Document) -> None:
    """Strip borders from every table in the document."""

    for table in document.tables:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "nil")
            borders.append(border)
        tbl_pr.append(borders)


def apply_yahei_font(run, size_pt: float) -> None:
    """Give a run a Microsoft YaHei font face and size."""

    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size_pt)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr, font in (
        ("w:ascii", "Microsoft YaHei"),
        ("w:hAnsi", "Microsoft YaHei"),
        ("w:cs", "Microsoft YaHei"),
        ("w:eastAsia", "微软雅黑"),
    ):
        r_fonts.set(qn(attr), font)


def format_cell_text(cell, font_size: float) -> None:
    """Ensure every run in the cell uses the right font and size."""

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            apply_yahei_font(run, font_size)


def strip_leading_number(title: str) -> str:
    """Drop existing numeric prefixes so Word numbering is fresh."""

    return re.sub(r"^\s*\d+[\.)]\s*", "", title)


def build_tables(items: Iterable[dict], document: Document) -> None:
    """Add a table per item to the document."""

    for item in items:
        table = document.add_table(rows=3, cols=2)
        table.allow_autofit = False

        # merge the first row horizontally for the title
        title_cell = table.cell(0, 0)
        title_cell.merge(table.cell(0, 1))
        title_paragraph = title_cell.paragraphs[0]
        title_paragraph.style = "List Number"
        title_paragraph.text = strip_leading_number(item["title"].strip())
        format_cell_text(title_cell, 12)

        options = item["options"]
        letters: Sequence[str] = ["A", "B", "C", "D"]
        coords = [(1, 0), (1, 1), (2, 0), (2, 1)]

        for letter, coordinate in zip(letters, coords):
            cell = table.cell(*coordinate)
            value = options[letter.lower()].strip()
            cell.text = f"{letter}. {value}"
            format_cell_text(cell, 10)

        # spacing handled after formatting

        document.add_paragraph()  # add spacing between tables


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a borderless Word document from structured JSON."
    )
    parser.add_argument("input", type=Path, help="Path to the JSON file")
    parser.add_argument(
        "output", type=Path, help="Destination path for the Word document"
    )
    args = parser.parse_args()

    def ensure_docx_extension(path: Path) -> Path:
        if path.suffix.lower() == ".docx":
            return path
        return path.with_suffix(".docx")

    output_path = ensure_docx_extension(args.output)
    if output_path != args.output:
        print(
            f"Output file renamed to {output_path.name} to keep a valid .docx extension."
        )

    items = load_data(args.input)
    document = Document(output_path) if output_path.exists() else Document()
    build_tables(items, document)
    remove_table_borders(document)
    document.save(output_path)


if __name__ == "__main__":
    main()