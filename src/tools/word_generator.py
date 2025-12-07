"""
Word Generator Tool for LangChain Agent.

This tool generates Word documents from extracted questions,
supporting both multiple-choice and true/false question formats.
"""

import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from langchain.tools import tool


# ==================== Word Document Helpers ====================

def remove_table_borders(document: Document) -> None:
    """Remove borders from all tables in the document."""
    for table in document.tables:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "nil")
            borders.append(border)
        tbl_pr.append(borders)


def apply_yahei_font(run, size_pt: float) -> None:
    """Apply Microsoft YaHei font to a run."""
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size_pt)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr, font in (
        ("w:ascii", "Microsoft YaHei"),
        ("w:hAnsi", "Microsoft YaHei"),
        ("w:cs", "Microsoft YaHei"),
        ("w:eastAsia", "微软雅黑"),
    ):
        r_fonts.set(qn(attr), font)


def format_cell_text(cell, font_size: float) -> None:
    """Format text in a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            apply_yahei_font(run, font_size)


def strip_leading_number(title: str) -> str:
    """Remove leading numbers from question titles."""
    return re.sub(r"^\s*\d+[\.)、．]\s*", "", title)


# ==================== Multiple Choice Document Builder ====================

def build_multiple_choice_tables(items: Iterable[dict], document: Document) -> int:
    """Build tables for multiple choice questions.
    
    Args:
        items: List of question dictionaries
        document: Word document to add to
        
    Returns:
        Number of questions added
    """
    count = 0
    for item in items:
        table = document.add_table(rows=3, cols=2)
        table.allow_autofit = False

        # Merge first row for title
        title_cell = table.cell(0, 0)
        title_cell.merge(table.cell(0, 1))
        title_paragraph = title_cell.paragraphs[0]
        title_paragraph.style = "List Number"
        title_paragraph.text = strip_leading_number(item["title"].strip())
        format_cell_text(title_cell, 12)

        # Add options
        options = item.get("options", {})
        letters = ["A", "B", "C", "D"]
        coords = [(1, 0), (1, 1), (2, 0), (2, 1)]

        for letter, coordinate in zip(letters, coords):
            cell = table.cell(*coordinate)
            value = options.get(letter.lower(), "").strip()
            cell.text = f"{letter}. {value}"
            format_cell_text(cell, 10)

        document.add_paragraph()
        count += 1
    
    return count


# ==================== True/False Document Builder ====================

def build_true_false_paragraphs(items: Iterable[dict], document: Document) -> int:
    """Build paragraphs for true/false questions.
    
    Args:
        items: List of question dictionaries
        document: Word document to add to
        
    Returns:
        Number of questions added
    """
    count = 0
    for item in items:
        paragraph = document.add_paragraph()
        paragraph.style = "List Number"
        
        title_text = strip_leading_number(item["title"].strip())
        run = paragraph.add_run(f"{title_text}（    ）")
        apply_yahei_font(run, 12)
        
        count += 1
    
    return count


def build_true_false_tables(items: Iterable[dict], document: Document) -> int:
    """Build tables for true/false questions (for consistent formatting).
    
    Args:
        items: List of question dictionaries
        document: Word document to add to
        
    Returns:
        Number of questions added
    """
    count = 0
    for item in items:
        table = document.add_table(rows=1, cols=1)
        table.allow_autofit = False

        title_cell = table.cell(0, 0)
        title_paragraph = title_cell.paragraphs[0]
        title_paragraph.style = "List Number"
        
        title_text = strip_leading_number(item["title"].strip())
        title_paragraph.text = f"{title_text}（    ）"
        format_cell_text(title_cell, 12)

        document.add_paragraph()
        count += 1
    
    return count


# ==================== Section Management Helpers ====================

MC_HEADER_TEXT = "一、选择题"
TF_HEADER_TEXT = "二、判断题"

def find_paragraph_by_text(document: Document, text: str):
    """Find a paragraph containing specific text."""
    for p in document.paragraphs:
        if text in p.text:
            return p
    return None

def move_elements_to_before(elements: list, target_paragraph) -> None:
    """Move specified elements to before target_paragraph."""
    if not target_paragraph:
        return
        
    target_element = target_paragraph._element
    for element in elements:
        target_element.addprevious(element)

def insert_questions_section(
    document: Document,
    questions: list[dict],
    header_text: str,
    next_header_text: str | None,
    is_multiple_choice: bool,
    use_table: bool = True
) -> int:
    """Insert questions into the correct section."""
    header = find_paragraph_by_text(document, header_text)
    next_header = find_paragraph_by_text(document, next_header_text) if next_header_text else None
    
    insertion_point = None
    
    if header:
        # Header exists
        if next_header:
            insertion_point = next_header
        else:
            insertion_point = None
    else:
        # Header missing, create it
        if next_header:
            # Insert header before next_header
            header = next_header.insert_paragraph_before()
            run = header.add_run(header_text)
            apply_yahei_font(run, 14)
            run.bold = True
            insertion_point = next_header
        else:
            # Append header
            header = document.add_paragraph()
            run = header.add_run(header_text)
            apply_yahei_font(run, 14)
            run.bold = True
            insertion_point = None
            
    # Capture existing elements to identify new ones
    existing_elements = list(document.element.body)
    
    # Generate questions
    if is_multiple_choice:
        count = build_multiple_choice_tables(questions, document)
    else:
        if use_table:
            count = build_true_false_tables(questions, document)
        else:
            count = build_true_false_paragraphs(questions, document)
        
    # Identify new elements
    new_elements = [e for e in document.element.body if e not in existing_elements]
    
    # Move elements if needed
    if insertion_point:
        move_elements_to_before(new_elements, insertion_point)
        
    return count


# ==================== Main Generation Function ====================

def generate_word_document(
    questions: list[dict],
    output_path: Path,
    question_type: str,
    use_table: bool = True,
    append: bool = False
) -> tuple[bool, str, int]:
    """Generate a Word document from questions.
    
    Args:
        questions: List of question dictionaries
        output_path: Path to save the document
        question_type: "multiple_choice" or "true_false"
        use_table: Use table format for true/false questions
        append: Append to existing document if True
        
    Returns:
        Tuple of (success, message, questions_added)
    """
    try:
        # Ensure directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Load or create document
        if append and output_path.exists():
            document = Document(output_path)
        else:
            document = Document()
        
        # Build document based on question type
        if question_type == "multiple_choice":
            count = insert_questions_section(
                document, 
                questions, 
                MC_HEADER_TEXT, 
                TF_HEADER_TEXT, 
                is_multiple_choice=True
            )
            remove_table_borders(document)
        else:  # true_false
            count = insert_questions_section(
                document, 
                questions, 
                TF_HEADER_TEXT, 
                None, 
                is_multiple_choice=False,
                use_table=use_table
            )
            if use_table:
                remove_table_borders(document)
        
        # Save document
        document.save(output_path)
        
        return True, f"Successfully saved {count} questions to {output_path}", count
        
    except Exception as e:
        return False, f"Error generating document: {str(e)}", 0


# ==================== LangChain Tools ====================

@tool
def save_questions_word(
    questions_json: str,
    output_path: str,
    question_type: str = "auto",
    append: bool = False
) -> str:
    """Save questions to a Word document.
    
    This tool generates a formatted Word document from question data.
    Multiple choice questions are displayed in tables with options.
    True/false questions are displayed as numbered statements.
    
    Args:
        questions_json: JSON string containing questions. Supports two formats:
                       1. Array format: [{\"title\": \"...\", \"options\": {...}}] or [{\"title\": \"...\"}]
                       2. Mixed format from image analysis: {\"multiple_choice\": [...], \"true_false\": [...]}
        output_path: File path where the Word document will be saved.
                    Will add .docx extension if not present.
        question_type: Type of questions to process. Options:
                      - \"auto\": Auto-detect from data structure, or process all if mixed
                      - \"multiple_choice\": Process as multiple choice (tables with 4 options)
                      - \"true_false\": Process as true/false (numbered statements)
        append: If True, append to existing document.
               If False, create new document. Default: False
    
    Returns:
        A string describing the result of the operation.
    """
    # Parse input questions
    try:
        data = json.loads(questions_json)
    except json.JSONDecodeError as e:
        return f"Error: Invalid JSON: {str(e)}"
    
    # Handle mixed format from image analysis
    if isinstance(data, dict) and ("multiple_choice" in data or "true_false" in data):
        mc_questions = data.get("multiple_choice", [])
        tf_questions = data.get("true_false", [])
        
        if question_type == "multiple_choice":
            questions = mc_questions
            effective_type = "multiple_choice"
        elif question_type == "true_false":
            questions = tf_questions
            effective_type = "true_false"
        else:  # auto - process both types
            # Convert to Path object
            file_path = Path(output_path)
            if file_path.suffix.lower() != ".docx":
                file_path = file_path.with_suffix(".docx")
            
            # Generate document with both types
            return _generate_mixed_word_document(
                mc_questions, tf_questions, file_path, append
            )
    elif isinstance(data, list):
        questions = data
        # Auto-detect type from data structure
        if question_type == "auto":
            # Check if questions have options (multiple choice) or not (true/false)
            if questions and "options" in questions[0]:
                effective_type = "multiple_choice"
            else:
                effective_type = "true_false"
        else:
            effective_type = question_type
    else:
        return "Error: Questions must be a JSON array or mixed format object"
    
    if not questions:
        return "Error: No questions provided. The questions array is empty."
    
    # Validate question type
    if effective_type not in ("multiple_choice", "true_false"):
        return f"Error: Invalid question_type '{effective_type}'. Must be 'multiple_choice', 'true_false', or 'auto'."
    
    # Validate question format
    if effective_type == "multiple_choice":
        for i, q in enumerate(questions):
            if "title" not in q:
                return f"Error: Question {i+1} is missing 'title' field"
            if "options" not in q:
                return f"Error: Question {i+1} is missing 'options' field"
    else:
        for i, q in enumerate(questions):
            if "title" not in q:
                return f"Error: Question {i+1} is missing 'title' field"
    
    # Convert to Path object
    file_path = Path(output_path)
    
    # Ensure .docx extension
    if file_path.suffix.lower() != ".docx":
        file_path = file_path.with_suffix(".docx")
    
    # Generate document
    success, message, count = generate_word_document(
        questions=questions,
        output_path=file_path,
        question_type=effective_type,
        use_table=True,
        append=append
    )
    
    if success:
        mode = "Appended to" if append and file_path.exists() else "Created"
        return f"{mode} Word document: {file_path}\n{message}"
    
    return message


def _generate_mixed_word_document(
    mc_questions: list[dict],
    tf_questions: list[dict],
    output_path: Path,
    append: bool = False
) -> str:
    """Generate a Word document with both multiple choice and true/false questions.
    
    Args:
        mc_questions: List of multiple choice question dictionaries
        tf_questions: List of true/false question dictionaries
        output_path: Path to save the document
        append: Append to existing document if True
        
    Returns:
        Result message string
    """
    try:
        # Ensure directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Load or create document
        if append and output_path.exists():
            document = Document(output_path)
        else:
            document = Document()
        
        mc_count = 0
        tf_count = 0
        
        # Add multiple choice questions first
        if mc_questions:
            mc_count = insert_questions_section(
                document, 
                mc_questions, 
                MC_HEADER_TEXT, 
                TF_HEADER_TEXT, 
                is_multiple_choice=True
            )
        
        # Add true/false questions
        if tf_questions:
            tf_count = insert_questions_section(
                document, 
                tf_questions, 
                TF_HEADER_TEXT, 
                None, 
                is_multiple_choice=False,
                use_table=True
            )
        
        # Remove all table borders
        remove_table_borders(document)
        
        # Save document
        document.save(output_path)
        
        total_count = mc_count + tf_count
        mode = "Appended to" if append and output_path.exists() else "Created"
        return f"{mode} Word document: {output_path}\nSuccessfully saved {total_count} questions ({mc_count} multiple choice, {tf_count} true/false)"
        
    except Exception as e:
        return f"Error generating document: {str(e)}"


# Export for convenient access
__all__ = [
    "save_questions_word",
    "generate_word_document",
    "_generate_mixed_word_document",
]
